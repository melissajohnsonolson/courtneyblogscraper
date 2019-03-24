# -*- coding: utf-8 -*-
"""
Created on Tue Feb  5 20:08:00 2019

@author: johns
"""

from __future__ import print_function
from bs4 import BeautifulSoup
import requests
from six.moves.urllib import parse
import pandas as pd
from docx import Document
import re
import os

#This is the location where you want to store the final word document
os.chdir('C:/mydirectory...')

#Here is the start page, or first page of the blog
start_page = "https://redhogdiary.wordpress.com/"

#This is a list that temporarily stores the url address for the next page
#before it's called for scrapping
queue = []

#All of the urls are stored here. This also ensures the script stops at the
#end of the blog, rather
#than alternately loading the last two pages infinitely
list_links = []

#The title for each blog post is stored here
titles = []
#The text for each blog post is stored here
entries = []

#this function connects to a url and scrapes the page for blog entry titles,
#text, and the link for the next page.
def link_getter(url):
    #we start by connecting to the webpage and using an html parser
    #to navigate the page
    r = requests.get(url)
    soup = BeautifulSoup(r.text, "html.parser")
    #All of the blog titles have an 'a' tag and have no class,
    #but are marked as a title.  This findAll collects all of the elements
    #that match the tag and attribute
    titles_group = soup.findAll('a', attrs = {'title':True, 'class':False})
    #Now we loop through the tagged elements to extract the title text
    for i in range(1, len(titles_group)):
         titles.append(titles_group[i]['title'])
    #All of the blog entries are in a div element with a class of 'entry' We
    #first collect all of those elements on the page, then loop through them
    #to grab the individual post text and store it in the entries list.
    entries_group = soup.findAll('div', {'class':'entry'})
    for i in range(0, len(entries_group)):
         entries.append(entries_group[i].text)
    
    #Here is where we find the url for the next page.  It falls under its own
    #div block with a class name of navigation. Typically the last entry in this
    #set of  results is the next page, until we get to the last page, 
    #where the last entry is the previous page's url.  This is why we check
    #list_links list to make sure we haven't accessed it before
    links = soup.find('div', {'class':'navigation'}).findAll('a')
    next_page=links[-1]['href']
    if next_page not in list_links:
        #if we have not accessed the link before, we append it to the queue
        #and to the list of urls. We also append the name of the this function
        #to the queue as we'll use it to run the function as we run through
        #items in the queue. We print out the progress so we can monitor
        #the program.
        queue.append((link_getter, next_page))
        list_links.append(next_page)
        print("processing " + next_page)

#This is the function that runs everything.  It calls the first link and 
#function that's currently in the queue 
def main():
    #we start by putting the starting page and the link_getter function into
    #the queue
    queue.append((link_getter, start_page))
    while len(queue):
        #as long as there is an item in the queue, this part runs. We call
        #the function name and the url from the top of the queue and then delete
        #it (that's what the .pop() method does)
        func, url = queue.pop(0)
        #func was the name of our function, and we pass it the url so it will
        #run.
        func(url)
        
main()

#The next step is to neatly write everything into a word document.  The end of
#each entry has some extra text we don't really want, like "ADVERTISEMENT'
#and "Comments" and such. These seem to always appear after a double line break.
#So we create a regular expression pattern to grab all the text before a 
#a double line break.  This pattern is stored as sequence.
sequence = re.compile(r"^(.+?)\n\n", re.DOTALL)

#this call starts a new document
document = Document()

#We loop through the titles and entries list. Double check that they have
#the same number of entires
for i in range(0, len(titles)):
     #Here we put the title into the word doc as a level 1 heading
    title_text = str(titles[i])
    document.add_heading(title_text, level =1)
    
    #Here we collect all the text before the double line break then add it as
    #a paragraph
    entry = sequence.search(entries[i]).group(0)
    paragraph = document.add_paragraph()
    paragraph.add_run(entry)

#Here we save our word document with whatever file name we choose.    
document.save('red hog diaries.docx')
