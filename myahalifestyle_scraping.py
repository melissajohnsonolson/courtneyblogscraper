# -*- coding: utf-8 -*-
"""
Created on Tue Jan 29 20:04:22 2019

@author: johns
"""

import requests
import urllib.request
import time, os
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

os.chdir('C:/mydirectory...')

#To start we want to create a list of links to every post.
#We start with the main url and then tack on page numbers
#(the first for loop). By inspecting the homepage, we see there are 5
#separate sub-pages. Each sub-page displays previews of several blog posts.
#To get the whole blog post, we need to go to each individual post's
#page. The link to the post's page can be found by searching for all
#h2 with a class of entry-title.  We loop through that result for each 
#sub-page, grabbing the link.
url = 'http://myahalifestyle.com/index.php/page/'
response = requests.get(url)
post_urls = []
for i in range(1,5):
    #this connects us to a specific page
    response = requests.get(url+str(i))
    page = BeautifulSoup(response.text, 'html.parser' )
    links = page.findAll('h2', {'class':'entry-title'})
    for j in range(0, len(links)):
        post=links[j].find('a')
        link = post['href']
        post_urls.append(link)
        
        


#Now that we have a list of all of the links, we call each one
#in the loop. We grab the title and the text, then
#append them to our dataframe
posts=pd.DataFrame(columns = ['URL', 'Title', 'Post'])
for i in range(0,len(post_urls)):
    site = post_urls[i]
    newresponse = requests.get(site)
    post = BeautifulSoup(newresponse.text, 'html.parser')
    #We grab the title of the post...
    title = post.findAll(class_ = 'entry-title')
    title = title[0].text.strip()
    #And we grab the text of the post....
    text = post.findAll(class_ = 'entry-content')
    text = text[0].text.strip()
    #There are some strange line breaks in these posts. Eventually
    #I would like every line break to be a double for appear, 
    #so I started by converting all multi-line breaks into single
    #breaks, then convterted those to a double line break.
    text = text.replace('\n\n\n', '\n')
    text = text.replace('\n\n', '\n')
    text = text.replace('\n', '\n\n')
    new_entry = [site, title, text]
    posts.loc[i] = new_entry

#Now we can creat the word document. We start by initializing a document
#then looping through the dataframe to add titles and posts in different
#formats
document = Document()
for i in range(0, len(posts['Post Number'])):
     
    title_text = str(posts.Title[i])
    document.add_heading(title_text, level =1)
    
    paragraph = document.add_paragraph()
    paragraph.add_run(posts.Post[i])
    
document.save('my aha lifestyle.docx')
